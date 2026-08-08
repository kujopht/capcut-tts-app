"""
Nhap van ban tu file .txt / .md / .docx va tu thu muc.

Nguyen tac:
- KHONG BAO GIO ghi/sua file nguon (chi mo che do doc).
- Ho tro UTF-8 va Unicode tieng Viet; ten file/duong dan co dau va khoang trang.
- File rong hoac khong doc duoc -> tra ve InputItem co `error`, khong nem
  exception ra ngoai, khong lam dong ung dung.
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable, List, Optional, Sequence, Tuple

from desktop_app.models import (
    SUPPORTED_EXTENSIONS,
    ErrorKind,
    InputItem,
    InputKind,
)

#: Cac bang ma thu lan luot khi file khong phai UTF-8 hop le
_ENCODING_FALLBACKS: Tuple[str, ...] = ("utf-8-sig", "utf-8", "utf-16", "cp1258", "cp1252", "latin-1")

FILE_DIALOG_FILTER = (
    "Văn bản hỗ trợ (*.txt *.md *.docx);;"
    "Text (*.txt);;Markdown (*.md);;Word (*.docx);;Tất cả file (*)"
)


def is_supported(path: Path | str) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS


def _read_plain_text(path: Path) -> str:
    """Doc .txt/.md, thu lan luot cac bang ma. Chi mo che do doc nhi phan."""
    data = path.read_bytes()
    if not data.strip():
        return ""
    last_error: Optional[Exception] = None
    for encoding in _ENCODING_FALLBACKS:
        try:
            text = data.decode(encoding)
        except (UnicodeDecodeError, LookupError) as exc:
            last_error = exc
            continue
        # utf-16 chi dung khi that su co BOM, tranh giai ma sai thanh rac
        if encoding == "utf-16" and not (data[:2] in (b"\xff\xfe", b"\xfe\xff")):
            continue
        return text.replace("\r\n", "\n").replace("\r", "\n")
    raise UnicodeDecodeError(
        "utf-8", data[:32], 0, 1,
        f"không giải mã được bằng {', '.join(_ENCODING_FALLBACKS)} ({last_error})",
    )


def _read_docx(path: Path) -> str:
    """
    Doc .docx bang python-docx. Lay text cua tung paragraph, giu ranh gioi doan,
    va lay ca text trong bang (table) neu co.
    """
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - phu thuoc moi truong
        raise RuntimeError(
            "Thiếu package 'python-docx' nên không đọc được .docx. "
            "Chạy: pip install -r requirements-gui.txt"
        ) from exc

    document = docx.Document(str(path))
    blocks: List[str] = []
    for paragraph in document.paragraphs:
        blocks.append(paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))

    # Gop cac dong, giu 1 dong trong giua cac doan co noi dung
    lines: List[str] = []
    for block in blocks:
        if block:
            lines.append(block)
        elif lines and lines[-1] != "":
            lines.append("")
    return "\n".join(lines).strip()


def read_text_file(path: Path | str) -> str:
    """
    Doc noi dung van ban tu mot file duoc ho tro.
    Nem exception khi that bai (ham goi phia tren se bat va danh dau loi).
    """
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix == ".docx":
        return _read_docx(file_path)
    if suffix in (".txt", ".md"):
        return _read_plain_text(file_path)
    # Duoi khong ho tro: van thu doc dang text de nguoi dung khong bi chan vo co
    return _read_plain_text(file_path)


def import_file(path: Path | str) -> InputItem:
    """
    Tao InputItem tu mot file. Luon tra ve InputItem — loi duoc ghi vao
    `item.error` / `item.error_kind` thay vi nem ra ngoai.
    """
    file_path = Path(path)
    item = InputItem(
        name=file_path.stem or file_path.name,
        kind=InputKind.FILE,
        path=str(file_path),
    )

    try:
        if not file_path.exists():
            item.error = "File không tồn tại"
            item.error_kind = ErrorKind.READ_FILE_ERROR.value
            return item
        if not file_path.is_file():
            item.error = "Đường dẫn không phải file"
            item.error_kind = ErrorKind.READ_FILE_ERROR.value
            return item
        if file_path.stat().st_size == 0:
            item.error = "File rỗng (0 byte)"
            item.error_kind = ErrorKind.EMPTY_TEXT.value
            return item

        text = read_text_file(file_path)
        if not (text or "").strip():
            item.error = "File không có nội dung văn bản"
            item.error_kind = ErrorKind.EMPTY_TEXT.value
            return item
        item.text = text
    except UnicodeDecodeError as exc:
        item.error = f"Không giải mã được nội dung: {exc.reason if hasattr(exc, 'reason') else exc}"
        item.error_kind = ErrorKind.READ_FILE_ERROR.value
    except PermissionError:
        item.error = "Không có quyền đọc file (file có thể đang được mở)"
        item.error_kind = ErrorKind.READ_FILE_ERROR.value
    except RuntimeError as exc:
        item.error = str(exc)
        item.error_kind = ErrorKind.READ_FILE_ERROR.value
    except Exception as exc:  # khong de bat ky loi nao lam dong app
        item.error = f"{type(exc).__name__}: {exc}"
        item.error_kind = ErrorKind.READ_FILE_ERROR.value

    return item


def import_files(paths: Iterable[Path | str]) -> List[InputItem]:
    """Nhap nhieu file trong mot lan. File loi van co mat trong danh sach."""
    return [import_file(path) for path in paths]


def collect_supported_files(
    directory: Path | str,
    recursive: bool = True,
) -> List[Path]:
    """
    Liet ke toan bo file duoc ho tro trong mot thu muc.
    Bo qua file tam cua Word (~$...) va thu muc an cua he thong.
    """
    root = Path(directory)
    if not root.is_dir():
        return []
    pattern = "**/*" if recursive else "*"
    found: List[Path] = []
    try:
        for path in sorted(root.glob(pattern)):
            if not path.is_file():
                continue
            if path.name.startswith("~$") or path.name.startswith("."):
                continue
            if is_supported(path):
                found.append(path)
    except OSError:
        return found
    return found


def import_directory(directory: Path | str, recursive: bool = True) -> List[InputItem]:
    """Nhap toan bo file duoc ho tro trong mot thu muc."""
    return import_files(collect_supported_files(directory, recursive=recursive))


def import_paths(paths: Sequence[Path | str], recursive: bool = True) -> List[InputItem]:
    """
    Nhap tu danh sach duong dan bat ky (dung cho keo-tha): file thi nhap file,
    thu muc thi nhap toan bo file duoc ho tro ben trong.
    """
    items: List[InputItem] = []
    for raw in paths:
        path = Path(raw)
        if path.is_dir():
            items.extend(import_directory(path, recursive=recursive))
        else:
            items.append(import_file(path))
    return items


def make_text_item(text: str, name: str = "van_ban_nhap_truc_tiep") -> InputItem:
    """Tao InputItem tu van ban nguoi dung go/dan truc tiep."""
    item = InputItem(name=name, kind=InputKind.TEXT, text=text or "")
    if not (text or "").strip():
        item.error = "Chưa nhập nội dung"
        item.error_kind = ErrorKind.EMPTY_TEXT.value
    return item
