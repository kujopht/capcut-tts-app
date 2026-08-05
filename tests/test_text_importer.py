"""Test nhap van ban tu .txt / .md / .docx, Unicode tieng Viet, va cac ca loi."""

from __future__ import annotations

import tempfile
import unittest
from pathlib import Path

from desktop_app.models import ErrorKind, InputKind
from desktop_app.text_importer import (
    collect_supported_files,
    import_directory,
    import_file,
    import_files,
    import_paths,
    is_supported,
    make_text_item,
    read_text_file,
)

VIETNAMESE = (
    "Chương một: Đường về nhà.\n\n"
    "Cô ấy nói: “Anh có nghe thấy tiếng mưa không?”\n"
    "Đứa trẻ mỉm cười — rồi lặng lẽ quay đi. Ừ, đã muộn rồi…\n"
)


class TestPlainText(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.TemporaryDirectory()
        self.dir = Path(self.tmp.name)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_txt_utf8_vietnamese(self) -> None:
        path = self.dir / "chuong_01.txt"
        path.write_text(VIETNAMESE, encoding="utf-8")
        item = import_file(path)
        self.assertFalse(item.error, item.error)
        self.assertIn("Đường về nhà", item.text)
        self.assertIn("“Anh có nghe thấy tiếng mưa không?”", item.text)
        self.assertEqual(item.kind, InputKind.FILE)
        self.assertEqual(item.name, "chuong_01")
        self.assertTrue(item.is_valid)

    def test_md_file(self) -> None:
        path = self.dir / "note.md"
        path.write_text("# Tiêu đề\n\nNội dung **đậm** tiếng Việt.\n", encoding="utf-8")
        item = import_file(path)
        self.assertFalse(item.error)
        self.assertIn("Tiêu đề", item.text)

    def test_utf8_bom_is_stripped(self) -> None:
        path = self.dir / "bom.txt"
        path.write_text(VIETNAMESE, encoding="utf-8-sig")
        item = import_file(path)
        self.assertFalse(item.error)
        self.assertTrue(item.text.startswith("Chương"), item.text[:20])

    def test_utf16_with_bom(self) -> None:
        path = self.dir / "u16.txt"
        path.write_text(VIETNAMESE, encoding="utf-16")
        item = import_file(path)
        self.assertFalse(item.error, item.error)
        self.assertIn("Đường về nhà", item.text)

    def test_crlf_normalized(self) -> None:
        path = self.dir / "crlf.txt"
        path.write_bytes("Dòng 1\r\nDòng 2\r\n".encode("utf-8"))
        item = import_file(path)
        self.assertNotIn("\r", item.text)
        self.assertIn("Dòng 1\nDòng 2", item.text)

    def test_filename_with_spaces_and_diacritics(self) -> None:
        path = self.dir / "Chương 1 — Đường về nhà (bản nháp).txt"
        path.write_text(VIETNAMESE, encoding="utf-8")
        item = import_file(path)
        self.assertFalse(item.error, item.error)
        self.assertIn("Đường về nhà", item.text)
        # slug dung lam ten thu muc phai an toan
        self.assertNotIn(" ", item.slug)
        self.assertTrue(item.slug.startswith("chuong_1"))

    def test_source_file_not_modified(self) -> None:
        path = self.dir / "readonly.txt"
        path.write_text(VIETNAMESE, encoding="utf-8")
        before = (path.read_bytes(), path.stat().st_mtime_ns)
        import_file(path)
        after = (path.read_bytes(), path.stat().st_mtime_ns)
        self.assertEqual(before, after, "File nguồn không được phép bị sửa")

    def test_empty_file_is_marked_error(self) -> None:
        path = self.dir / "empty.txt"
        path.write_text("", encoding="utf-8")
        item = import_file(path)
        self.assertTrue(item.error)
        self.assertEqual(item.error_kind, ErrorKind.EMPTY_TEXT.value)
        self.assertFalse(item.is_valid)

    def test_whitespace_only_file_is_error(self) -> None:
        path = self.dir / "blank.txt"
        path.write_text("   \n\n\t\n", encoding="utf-8")
        item = import_file(path)
        self.assertTrue(item.error)
        self.assertFalse(item.is_valid)

    def test_missing_file_does_not_raise(self) -> None:
        item = import_file(self.dir / "khong_co.txt")
        self.assertTrue(item.error)
        self.assertEqual(item.error_kind, ErrorKind.READ_FILE_ERROR.value)

    def test_directory_passed_as_file(self) -> None:
        item = import_file(self.dir)
        self.assertTrue(item.error)

    def test_binary_garbage_does_not_crash(self) -> None:
        path = self.dir / "junk.txt"
        path.write_bytes(bytes(range(0, 255)) * 4)
        item = import_file(path)          # co the doc duoc bang latin-1, mien la khong crash
        self.assertIsNotNone(item)


class TestDocx(unittest.TestCase):
    def setUp(self) -> None:
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("python-docx chưa được cài")
        self.tmp = tempfile.TemporaryDirectory()
        self.dir = Path(self.tmp.name)

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def _make_docx(self, name: str = "truyen.docx") -> Path:
        import docx

        document = docx.Document()
        document.add_paragraph("Chương một: Đường về nhà")
        document.add_paragraph("")
        document.add_paragraph("Cô ấy nói: “Anh có nghe thấy tiếng mưa không?”")
        document.add_paragraph("Đứa trẻ mỉm cười — rồi lặng lẽ quay đi.")
        table = document.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "Nhân vật"
        table.cell(0, 1).text = "Hà Nội"
        path = self.dir / name
        document.save(str(path))
        return path

    def test_docx_vietnamese(self) -> None:
        path = self._make_docx()
        item = import_file(path)
        self.assertFalse(item.error, item.error)
        self.assertIn("Chương một: Đường về nhà", item.text)
        self.assertIn("“Anh có nghe thấy tiếng mưa không?”", item.text)
        self.assertIn("Đứa trẻ mỉm cười", item.text)

    def test_docx_table_text_included(self) -> None:
        item = import_file(self._make_docx())
        self.assertIn("Nhân vật", item.text)
        self.assertIn("Hà Nội", item.text)

    def test_docx_paragraph_breaks_preserved(self) -> None:
        item = import_file(self._make_docx())
        self.assertIn("\n", item.text)

    def test_empty_docx_is_error(self) -> None:
        import docx

        path = self.dir / "rong.docx"
        docx.Document().save(str(path))
        item = import_file(path)
        self.assertTrue(item.error)
        self.assertEqual(item.error_kind, ErrorKind.EMPTY_TEXT.value)

    def test_corrupt_docx_does_not_crash(self) -> None:
        path = self.dir / "hong.docx"
        path.write_bytes(b"day khong phai file docx")
        item = import_file(path)
        self.assertTrue(item.error)
        self.assertEqual(item.error_kind, ErrorKind.READ_FILE_ERROR.value)

    def test_docx_source_not_modified(self) -> None:
        path = self._make_docx()
        before = path.read_bytes()
        import_file(path)
        self.assertEqual(before, path.read_bytes())


class TestBatchImport(unittest.TestCase):
    def setUp(self) -> None:
        self.tmp = tempfile.TemporaryDirectory()
        self.dir = Path(self.tmp.name)
        (self.dir / "a.txt").write_text("Nội dung A", encoding="utf-8")
        (self.dir / "b.md").write_text("Nội dung B", encoding="utf-8")
        (self.dir / "c.pdf").write_text("không hỗ trợ", encoding="utf-8")
        (self.dir / "~$tam.docx").write_text("file tạm của Word", encoding="utf-8")
        sub = self.dir / "Thư mục con"
        sub.mkdir()
        (sub / "d.txt").write_text("Nội dung D", encoding="utf-8")

    def tearDown(self) -> None:
        self.tmp.cleanup()

    def test_is_supported(self) -> None:
        self.assertTrue(is_supported("x.txt"))
        self.assertTrue(is_supported("x.MD"))
        self.assertTrue(is_supported("x.docx"))
        self.assertFalse(is_supported("x.pdf"))

    def test_collect_skips_unsupported_and_temp(self) -> None:
        found = {p.name for p in collect_supported_files(self.dir)}
        self.assertEqual(found, {"a.txt", "b.md", "d.txt"})

    def test_collect_non_recursive(self) -> None:
        found = {p.name for p in collect_supported_files(self.dir, recursive=False)}
        self.assertEqual(found, {"a.txt", "b.md"})

    def test_import_directory(self) -> None:
        items = import_directory(self.dir)
        self.assertEqual(len(items), 3)
        self.assertTrue(all(i.is_valid for i in items))

    def test_import_multiple_files_in_one_go(self) -> None:
        items = import_files([self.dir / "a.txt", self.dir / "b.md"])
        self.assertEqual(len(items), 2)

    def test_import_paths_mixes_files_and_dirs(self) -> None:
        items = import_paths([self.dir / "a.txt", self.dir / "Thư mục con"])
        names = sorted(i.name for i in items)
        self.assertEqual(names, ["a", "d"])

    def test_bad_file_in_batch_does_not_stop_others(self) -> None:
        items = import_files([self.dir / "khong_co.txt", self.dir / "a.txt"])
        self.assertEqual(len(items), 2)
        self.assertTrue(items[0].error)
        self.assertFalse(items[1].error)
        self.assertTrue(items[1].is_valid)


class TestDirectText(unittest.TestCase):
    def test_make_text_item(self) -> None:
        item = make_text_item(VIETNAMESE, name="truyen_moi")
        self.assertEqual(item.kind, InputKind.TEXT)
        self.assertTrue(item.is_valid)
        self.assertEqual(item.source_label(), "(nhập trực tiếp)")
        self.assertEqual(item.char_count, len(VIETNAMESE))

    def test_make_text_item_empty_is_error(self) -> None:
        item = make_text_item("   ")
        self.assertTrue(item.error)
        self.assertFalse(item.is_valid)

    def test_read_text_file_direct(self) -> None:
        with tempfile.TemporaryDirectory() as tmp:
            path = Path(tmp) / "x.txt"
            path.write_text("Xin chào", encoding="utf-8")
            self.assertEqual(read_text_file(path), "Xin chào")


if __name__ == "__main__":
    unittest.main(verbosity=2)
